"""Relleno automático de la planilla oficial de inscripción (mail merge sobre .docx).

No genera un documento nuevo: abre `templates_docx/planilla_preinscripcion.docx`
tal cual (mismo encabezado, logo y tablas) y solo inserta los datos del alumno
en las celdas correspondientes, localizadas por el texto de su etiqueta
(ej. "APELLIDOS:"). La plantilla tiene 3 tablas, una por sección de la
planilla física (Estudiante / Representante / Datos Administrativos), en ese
orden — por eso cada tabla se procesa con su propio mapa de etiquetas en vez
de uno solo, ya que etiquetas como "APELLIDOS:" o "CÉDULA:" se repiten entre
secciones.
"""
from datetime import date
from io import BytesIO
from itertools import count
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

from cobranza.models import Pago
from cobranza.services import calcular_datos_administrativos_inscripcion
from .models import Representante

TEMPLATE_PATH = Path(__file__).resolve().parent / 'templates_docx' / 'planilla_preinscripcion.docx'

LABEL_MAP_ESTUDIANTE = {
    'APELLIDOS': 'apellido_estudiante',
    'NOMBRES': 'nombre_estudiante',
    'FECHA DE NACIMIENTO': 'fecha_nacimiento',
    'LUGAR DE NACIMIENTO': 'lugar_nacimiento',
    'PAÍS': 'pais_nacimiento',
    'ESTADO': 'estado_nacimiento',
    'CÉDULA': 'cedula_estudiante',
    'SEXO': 'sexo',
    'PESO': 'peso',
    'ESTATURA': 'estatura',
    'DIRECCION DE HABITACIÓN': 'direccion_estudiante',
    'EDAD': 'edad',
    'INSTITUCIÓN DE PROCEDENCIA': 'institucion_procedencia',
    'BAUTIZADO': 'bautizado',
    'CURSARÁ': 'cursara',
    'ALÉRGICO': 'alergico',
}

LABEL_MAP_REPRESENTANTE = {
    'APELLIDOS': 'apellido_representante',
    'NOMBRES': 'nombre_representante',
    'CÉDULA': 'cedula_representante',
    'PARENTESCO': 'parentesco',
    'DIRECCION DE HABITACIÓN': 'direccion_representante',
    'NACIONALIDAD': 'nacionalidad',
    'Nº TELÉFONO': 'telefono',
    'NIVEL DE ESTUDIO': 'nivel_estudio',
}

LABEL_MAP_ADMINISTRATIVO = {
    'Nº SOLVENCIA': 'nro_solvencia',
    'Nº DE TRANSFERENCIA': 'nro_transferencia',
    'TRASFERENCIA MONTO': 'monto_transferencia',
    'EFECTIVO MONTO': 'monto_efectivo',
    'BANCO DE PROCEDENCIA': 'banco_procedencia',
    'BANCO DESTINO': 'banco_destino',
    'FECHA DE PAGO': 'fecha_pago',
    'FECHA INSCRIPCIÓN': 'fecha_inscripcion',
}

TODOS_LOS_CAMPOS = (
    set(LABEL_MAP_ESTUDIANTE.values())
    | set(LABEL_MAP_REPRESENTANTE.values())
    | set(LABEL_MAP_ADMINISTRATIVO.values())
)


def _normalizar_etiqueta(texto):
    texto = (texto or '').strip()
    if texto.endswith(':'):
        texto = texto[:-1]
    return ' '.join(texto.split()).upper()


_CONECTORES_MINUSCULA = {'de', 'del', 'la', 'las', 'los', 'y', 'e'}


def _capitalizar_nombre_propio(texto):
    """"Juan Perez" / "JUAN PEREZ" / "juan perez" -> "Juan Perez". Los conectores
    de apellidos compuestos ("de", "del", "la"...) quedan en minúscula salvo
    que abran el texto."""
    texto = (texto or '').strip()
    if not texto:
        return texto
    palabras = texto.lower().split()
    capitalizadas = [
        palabra if (i > 0 and palabra in _CONECTORES_MINUSCULA) else palabra[:1].upper() + palabra[1:]
        for i, palabra in enumerate(palabras)
    ]
    return ' '.join(capitalizadas)


def _celdas_unicas(table):
    """Las celdas fusionadas (gridSpan) comparten el mismo <w:tc> — sin este
    filtro se procesaría (y escribiría) la misma celda una vez por columna
    que abarca."""
    vistas = set()
    for row in table.rows:
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in vistas:
                continue
            vistas.add(tc_id)
            yield cell


def _escribir_valor(cell, valor):
    if not valor:
        return
    paragraphs = cell.paragraphs
    ultimo = paragraphs[-1]
    # Si la celda ya trae un párrafo en blanco reservado para la respuesta
    # (patrón de la mayoría de los campos), se escribe ahí; si no, se agrega
    # el valor a continuación de la etiqueta en el mismo renglón.
    if len(paragraphs) > 1 and not ultimo.text.strip():
        # Algunas etiquetas reservan hasta 2 párrafos en blanco antes del
        # que realmente recibe el valor (espacio pensado para llenado a
        # mano). En un campo autocompletado esos renglones vacíos de más
        # no cumplen ningún propósito y sólo agrandan la celda — se
        # eliminan, dejando únicamente la etiqueta y el párrafo con el
        # valor.
        for parrafo_vacio in paragraphs[1:-1]:
            if not parrafo_vacio.text.strip():
                parrafo_vacio._element.getparent().remove(parrafo_vacio._element)
        run = ultimo.add_run(str(valor))
    else:
        # Tab a un punto fijo del ancho de la celda en vez de 2 espacios
        # literales: con fuente proporcional, "APELLIDOS:" y "DIRECCION DE
        # HABITACIÓN:" ocupan un ancho visual muy distinto, así que un
        # separador de caracteres fijos no hace que el valor arranque en la
        # misma columna. Un tab stop sí, independientemente del largo de la
        # etiqueta. cell.width puede ser None si la plantilla no fija un
        # ancho explícito en esa celda (gridSpan poco común) — en ese caso
        # se cae al separador simple anterior.
        if cell.width:
            ultimo.paragraph_format.tab_stops.add_tab_stop(Emu(int(cell.width * 0.45)))
            run = ultimo.add_run(f'\t{valor}')
        else:
            run = ultimo.add_run(f'  {valor}')
    run.font.bold = False
    # 10pt (etiqueta queda en 8pt, w:sz=16 en la plantilla) para que el
    # valor se distinga de la etiqueta. Un valor más grande que la
    # etiqueta en la misma línea aumenta la altura de esa línea y,
    # multiplicado por las ~14 celdas que comparten renglón con su
    # etiqueta, es lo que puede desbordar a una segunda página — si eso
    # pasa, hay que revisar márgenes en vez de subir más este tamaño.
    run.font.size = Pt(10)


def _rellenar_tabla(table, mapa_etiquetas, valores):
    for cell in _celdas_unicas(table):
        if not cell.paragraphs:
            continue
        etiqueta = _normalizar_etiqueta(cell.paragraphs[0].text)
        clave = mapa_etiquetas.get(etiqueta)
        if not clave:
            continue
        _escribir_valor(cell, valores.get(clave, ''))


def _calcular_edad(fecha_nacimiento):
    hoy = date.today()
    return hoy.year - fecha_nacimiento.year - (
        (hoy.month, hoy.day) < (fecha_nacimiento.month, fecha_nacimiento.day)
    )


def _formatear_monto(monto):
    return f"$ {monto:.2f}" if monto else ''


_CODIGOS_TIPO_TRANSFERENCIA = ('transferencia', 'pago_movil', 'punto_de_venta', 'zelle')


def _combinar_metodos_transferencia(metodos_pago):
    """La planilla física solo tiene casillas de 'TRANSFERENCIA' y 'EFECTIVO'
    — sin campos propios para Pago Móvil, Punto de Venta o Zelle. Sin
    este agrupamiento esos métodos quedaban totalmente fuera del comprobante
    (ni monto ni referencia), aunque el pago sí existiera en el sistema.
    Decisión del usuario: tratarlos todos como 'transferencia' — se suman
    montos y referencias, y el banco de origen/destino se toma del primer
    método que sí tenga banco asociado (Pago Móvil/Punto de Venta/
    Transferencia; Zelle no tiene banco en `calcular_datos_administrativos_inscripcion`)."""
    displays = {dict(Pago.METODOS)[c] for c in _CODIGOS_TIPO_TRANSFERENCIA}
    grupos = [m for m in metodos_pago if m['metodo_display'] in displays]
    if not grupos:
        return None
    return {
        'monto': sum((g['monto'] for g in grupos), start=0),
        'referencia': ', '.join(g['referencia'] for g in grupos if g['referencia']),
        'banco_destino': next((g['banco_destino'] for g in grupos if g['banco_destino']), ''),
        'banco_procedencia': next((g['banco_procedencia'] for g in grupos if g['banco_procedencia']), ''),
    }


def _sumar_montos_efectivo(metodos_pago):
    displays = {dict(Pago.METODOS)['efectivo'], dict(Pago.METODOS)['efectivo_ves']}
    total = sum((m['monto'] for m in metodos_pago if m['metodo_display'] in displays), start=0)
    return total or None


def _resolver_valores(alumno, inscripcion, campos_seleccionados):
    seleccionados = set(campos_seleccionados) if campos_seleccionados else set(TODOS_LOS_CAMPOS)
    representante = alumno.representante

    if inscripcion:
        datos_admin = calcular_datos_administrativos_inscripcion(inscripcion)
    else:
        datos_admin = {'metodos_pago': [], 'fecha_pago': None, 'fecha_inscripcion': None, 'nro_solvencia': ''}

    transferencia = _combinar_metodos_transferencia(datos_admin['metodos_pago'])
    monto_efectivo = _sumar_montos_efectivo(datos_admin['metodos_pago'])

    crudos = {
        'apellido_estudiante': _capitalizar_nombre_propio(alumno.apellido),
        'nombre_estudiante': _capitalizar_nombre_propio(alumno.nombre),
        'fecha_nacimiento': alumno.fecha_nacimiento.strftime('%d/%m/%Y') if alumno.fecha_nacimiento else '',
        'lugar_nacimiento': alumno.lugar_nacimiento,
        'pais_nacimiento': alumno.pais_nacimiento,
        'estado_nacimiento': alumno.estado_nacimiento,
        # No existe cédula de identidad del alumno en el modelo (solo
        # cedula_escolar, que es matrícula interna) — decisión: queda en blanco.
        'cedula_estudiante': '',
        'sexo': alumno.get_genero_display() if alumno.genero else '',
        'peso': str(alumno.peso) if alumno.peso is not None else '',
        'estatura': str(alumno.estatura) if alumno.estatura is not None else '',
        'direccion_estudiante': _capitalizar_nombre_propio(alumno.direccion),
        'edad': str(_calcular_edad(alumno.fecha_nacimiento)) if alumno.fecha_nacimiento else '',
        'institucion_procedencia': _capitalizar_nombre_propio(alumno.institucion_procedencia),
        'bautizado': {True: 'Sí', False: 'No'}.get(alumno.bautizado, ''),
        'cursara': (inscripcion.grado_seccion if inscripcion else alumno.grado_seccion) or '',
        'alergico': alumno.alergico or '',

        'apellido_representante': _capitalizar_nombre_propio(representante.apellido) if representante else '',
        'nombre_representante': _capitalizar_nombre_propio(representante.nombre) if representante else '',
        'cedula_representante': representante.cedula if representante else '',
        'parentesco': dict(Representante.PARENTESCOS).get(alumno.parentesco, '') if alumno.parentesco else '',
        'direccion_representante': _capitalizar_nombre_propio(representante.direccion) if representante else '',
        'nacionalidad': representante.nacionalidad if representante else '',
        'telefono': representante.telefono if representante else '',
        'nivel_estudio': representante.nivel_estudio if representante else '',

        'nro_solvencia': datos_admin.get('nro_solvencia') or '',
        'nro_transferencia': transferencia['referencia'] if transferencia else '',
        'monto_transferencia': _formatear_monto(transferencia['monto']) if transferencia else '',
        'monto_efectivo': _formatear_monto(monto_efectivo),
        'banco_procedencia': transferencia['banco_procedencia'] if transferencia else '',
        'banco_destino': transferencia['banco_destino'] if transferencia else '',
        'fecha_pago': datos_admin['fecha_pago'].strftime('%d/%m/%Y') if datos_admin.get('fecha_pago') else '',
        'fecha_inscripcion': datos_admin['fecha_inscripcion'].strftime('%d/%m/%Y') if datos_admin.get('fecha_inscripcion') else '',
    }
    return {clave: (valor if clave in seleccionados else '') for clave, valor in crudos.items()}


def _generar_documento_preinscripcion(alumno, inscripcion, campos_seleccionados):
    """Devuelve el objeto Document (python-docx) de la planilla rellenada para
    un alumno, sin serializar — lo usan tanto la descarga individual como la
    combinación en un documento único."""
    valores = _resolver_valores(alumno, inscripcion, campos_seleccionados)

    doc = Document(str(TEMPLATE_PATH))
    _rellenar_tabla(doc.tables[0], LABEL_MAP_ESTUDIANTE, valores)
    _rellenar_tabla(doc.tables[1], LABEL_MAP_REPRESENTANTE, valores)
    _rellenar_tabla(doc.tables[2], LABEL_MAP_ADMINISTRATIVO, valores)
    return doc


def generar_planilla_preinscripcion(alumno, inscripcion, campos_seleccionados):
    """Devuelve un BytesIO con la planilla .docx rellenada para un alumno."""
    doc = _generar_documento_preinscripcion(alumno, inscripcion, campos_seleccionados)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def nombre_archivo_alumno(alumno):
    base = f"PreInscripcion_{alumno.apellido}_{alumno.nombre}_{alumno.id}".replace(' ', '_')
    return f"{base}.docx"


def generar_zip_preinscripciones(alumnos, campos_seleccionados):
    """Genera un .docx por alumno y los empaqueta en un .zip en memoria.

    Síncrono — sin cola de trabajo. Ver NOTAS_TECNICAS.md para el límite de
    este enfoque en datasets muy grandes.
    """
    from .models import Inscripcion

    zip_buffer = BytesIO()
    with ZipFile(zip_buffer, 'w') as zf:
        for alumno in alumnos:
            inscripcion = Inscripcion.objects.filter(alumno=alumno).order_by('-fecha_inscripcion').first()
            docx_buffer = generar_planilla_preinscripcion(alumno, inscripcion, campos_seleccionados)
            zf.writestr(nombre_archivo_alumno(alumno), docx_buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer


def _renumerar_docpr(doc, contador):
    """Reasigna w:id únicos a cada <wp:docPr> (identificador de cada imagen/
    dibujo) del documento, tomando valores de `contador` (un itertools.count
    compartido entre todos los documentos a combinar).

    La plantilla trae 3 imágenes en el cuerpo (logo, ver TEMPLATE_PATH), cada
    una con su propio docPr id (ej. id="9"). Como cada alumno se genera desde
    la misma plantilla, todos los documentos repiten esos mismos ids. Al
    combinarlos en un solo body, Word deja de poder resolver a qué imagen
    corresponde cada docPr id y considera el archivo corrupto al abrirlo
    ("El archivo parece estar corrompido") — confirmado reproduciendo el
    error con Word vía COM y aislándolo probando la combinación sin datos,
    solo con la plantilla repetida. Renumerar a valores únicos por documento
    lo evita; python-docx no valida esto (por eso el archivo se podía leer
    de vuelta con python-docx sin error, ocultando el problema hasta abrirlo
    en Word real)."""
    for elemento in doc.element.body.iter():
        if elemento.tag == qn('wp:docPr'):
            elemento.set('id', str(next(contador)))


def _insertar_salto_pagina_en_ultimo_parrafo(body):
    """Inserta un <w:br w:type="page"/> como run nuevo dentro del último <w:p>
    ya existente en `body`, en vez de usar Document.add_page_break() (que
    crea un <w:p> completamente nuevo para alojar el salto).

    La plantilla está ajustada al límite exacto de 1 página por planilla
    (ver "PLANTILLA — SEGUNDA PÁGINA EN BLANCO" en NOTAS_TECNICAS.md: margen
    inferior y tamaño de fuente recortados a propósito, sin margen de
    sobra). add_page_break() agrega un párrafo adicional que ya no cabe en
    esa página sin margen — ese párrafo desborda solo a la página
    siguiente, y como ES el que contiene el salto, termina generando una
    página en blanco antes de cada planilla siguiente (confirmado con Word
    vía COM: 3 documentos combinados con add_page_break() → 5 páginas en
    vez de 3, con un patrón exacto de [planilla][blanco][planilla][blanco]
    [planilla]). Reutilizar el último párrafo ya existente de la planilla
    anterior —que ya cabía en su página— evita agregar un párrafo nuevo:
    con este método, N documentos combinados dan exactamente N páginas."""
    ultimo_parrafo = None
    for elemento in reversed(body):
        if elemento.tag == qn('w:p'):
            ultimo_parrafo = elemento
            break
    if ultimo_parrafo is None:
        return
    run = OxmlElement('w:r')
    salto = OxmlElement('w:br')
    salto.set(qn('w:type'), 'page')
    run.append(salto)
    ultimo_parrafo.append(run)


def _combinar_documentos(documentos):
    """Concatena varios Document (python-docx) en uno solo, cada uno arrancando
    en página nueva. No usa docxcompose (no está en las dependencias del
    proyecto): copia los elementos del <w:body> de cada documento dentro del
    body del primero, salvo el <w:sectPr> final de cada uno (que define
    márgenes/tamaño de página y solo debe existir una vez, al final del
    documento combinado). Como todos parten de la misma plantilla, comparten
    header/footer/estilos — al ser un salto de página simple (no de sección),
    el encabezado con el logo se repite automáticamente en cada página."""
    if not documentos:
        raise ValueError("Sin documentos para combinar")

    contador_docpr = count(100000)
    for doc in documentos:
        _renumerar_docpr(doc, contador_docpr)

    maestro = documentos[0]
    body = maestro.element.body
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        body.remove(sect_pr)

    for doc in documentos[1:]:
        _insertar_salto_pagina_en_ultimo_parrafo(body)
        doc_body = doc.element.body
        doc_sect_pr = doc_body.find(qn('w:sectPr'))
        for elemento in list(doc_body):
            if elemento is doc_sect_pr:
                continue
            body.append(elemento)

    if sect_pr is not None:
        body.append(sect_pr)

    return maestro


def generar_documento_unico_preinscripciones(alumnos, campos_seleccionados):
    """Genera un solo .docx con la planilla de cada alumno activo, en orden
    alfabético por apellido/nombre, cada una en página nueva.

    Síncrono, igual que generar_zip_preinscripciones — ver NOTAS_TECNICAS.md
    para el límite de este enfoque en datasets muy grandes.
    """
    from .models import Inscripcion

    # .lower() explícito: apellido/nombre pueden venir en mayúsculas o
    # minúsculas mezcladas desde la BD, y una comparación case-sensitive
    # ordena todo lo que empieza en mayúscula antes que cualquier minúscula
    # (orden ASCII), no alfabéticamente como espera el usuario.
    alumnos_ordenados = sorted(alumnos, key=lambda a: ((a.apellido or '').lower(), (a.nombre or '').lower()))
    documentos = []
    for alumno in alumnos_ordenados:
        inscripcion = Inscripcion.objects.filter(alumno=alumno).order_by('-fecha_inscripcion').first()
        documentos.append(_generar_documento_preinscripcion(alumno, inscripcion, campos_seleccionados))

    maestro = _combinar_documentos(documentos)

    buffer = BytesIO()
    maestro.save(buffer)
    buffer.seek(0)
    return buffer
